import os
import json
from io import BytesIO

from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import Response
from pydantic import BaseModel
from sqlmodel import SQLModel, Session, create_engine, select
from anthropic import Anthropic
from docx import Document

from models import Message

DATABASE_URL = os.environ.get("DATABASE_URL", "sqlite:///./app.db")
if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}
engine = create_engine(DATABASE_URL, connect_args=connect_args)

app = FastAPI()
client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
MODEL = "claude-sonnet-4-6"

PROJECT_NAME = "Digital Commerce Project"

SYSTEM_PROMPT = f"""You are a project status assistant helping a project manager track the "{PROJECT_NAME}".

Your job is to have a focused conversation that surfaces the information needed to fill out periodic project status memos. Over time you want to learn about:
- The current phase and reporting period
- Schedule, scope, resources, and vendor health (green/yellow/red status, trend, notes)
- Recently completed work, work in progress, and upcoming work
- Risks (description, likelihood, impact, mitigation)
- Assumptions the team is operating under
- Open issues (who raised them, current status, next step)
- Decisions made (description, who made them, when)
- Anything else worth noting

Guidelines:
- Ask one or two specific follow-up questions at a time — don't interrogate.
- Reference details the user has already shared so they feel heard.
- Don't repeat questions already answered earlier in the conversation.
- If they share status colors, probe gently for the why and what's changing.
- When something sounds like a risk, issue, or decision, confirm the structured details (who, when, likelihood, impact, etc.).
- Be concise. Skip pleasantries and preamble."""


@app.on_event("startup")
def startup():
    SQLModel.metadata.create_all(engine)


class SendMessageRequest(BaseModel):
    content: str


def _serialize(m: Message) -> dict:
    return {
        "id": m.id,
        "role": m.role,
        "content": m.content,
        "timestamp": m.timestamp.isoformat(),
    }


@app.get("/api/messages")
def get_messages():
    with Session(engine) as session:
        messages = session.exec(select(Message).order_by(Message.timestamp)).all()
        return [_serialize(m) for m in messages]


@app.post("/api/messages")
def post_message(req: SendMessageRequest):
    content = req.content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    with Session(engine) as session:
        user_msg = Message(role="user", content=content)
        session.add(user_msg)
        session.commit()
        session.refresh(user_msg)

        all_messages = session.exec(select(Message).order_by(Message.timestamp)).all()
        conversation = [{"role": m.role, "content": m.content} for m in all_messages]

        response = client.messages.create(
            model=MODEL,
            max_tokens=1024,
            system=[
                {
                    "type": "text",
                    "text": SYSTEM_PROMPT,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            messages=conversation,
        )
        assistant_text = next(
            (b.text for b in response.content if b.type == "text"), ""
        ).strip()

        assistant_msg = Message(role="assistant", content=assistant_text)
        session.add(assistant_msg)
        session.commit()
        session.refresh(assistant_msg)

        return {
            "user_message": _serialize(user_msg),
            "assistant_message": _serialize(assistant_msg),
        }


MEMO_INSTRUCTIONS = f"""Based on the full conversation history with the project manager, produce a project status memo for the "{PROJECT_NAME}" as a single JSON object with EXACTLY this structure and these keys:

{{
  "report_date": "YYYY-MM-DD",
  "reporting_period_start": "YYYY-MM-DD",
  "reporting_period_end": "YYYY-MM-DD",
  "current_phase": "string",
  "prepared_by": "string",
  "schedule_status": "green | yellow | red",
  "schedule_trend": "improving | stable | declining",
  "schedule_notes": "string",
  "scope_status": "green | yellow | red",
  "scope_trend": "improving | stable | declining",
  "scope_notes": "string",
  "resources_status": "green | yellow | red",
  "resources_trend": "improving | stable | declining",
  "resources_notes": "string",
  "vendor_status": "green | yellow | red",
  "vendor_trend": "improving | stable | declining",
  "vendor_notes": "string",
  "completed_items": ["string"],
  "current_items": ["string"],
  "upcoming_items": ["string"],
  "risks": [
    {{"description": "string", "likelihood": "string", "impact": "string", "mitigation": "string"}}
  ],
  "assumptions": ["string"],
  "issues": [
    {{"description": "string", "raised_by": "string", "status": "string", "next_step": "string"}}
  ],
  "decisions": [
    {{"description": "string", "made_by": "string", "date": "YYYY-MM-DD"}}
  ],
  "notes": ["string"]
}}

Rules:
- Use only information present in the conversation. Do not invent specifics.
- For fields with no information yet, use empty strings, empty arrays, "green" for statuses with no concerns raised, and "stable" for trends.
- Today's date can be used for report_date if no explicit date was given.
- Return ONLY the JSON object, nothing else, no code fences, no commentary."""


def _strip_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines)
    return text.strip()


@app.post("/api/memo")
def generate_memo():
    with Session(engine) as session:
        messages = session.exec(select(Message).order_by(Message.timestamp)).all()

        if not messages:
            raise HTTPException(
                status_code=400, detail="No conversation history yet — chat first."
            )

        conversation_text = "\n\n".join(
            f"{m.role.upper()}: {m.content}" for m in messages
        )

        response = client.messages.create(
            model=MODEL,
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": (
                        f"{MEMO_INSTRUCTIONS}\n\n"
                        f"--- CONVERSATION ---\n{conversation_text}"
                    ),
                }
            ],
        )

        text = next((b.text for b in response.content if b.type == "text"), "")
        text = _strip_fences(text)

        try:
            memo = json.loads(text)
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse memo JSON from model: {e}",
            )

        return memo


class MemoDownloadRequest(BaseModel):
    memo: dict


def _add_kv(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value) if value not in (None, "") else "—")


def _add_status_block(doc, title, memo, prefix):
    doc.add_heading(title, level=2)
    _add_kv(doc, "Status", memo.get(f"{prefix}_status", ""))
    _add_kv(doc, "Trend", memo.get(f"{prefix}_trend", ""))
    _add_kv(doc, "Notes", memo.get(f"{prefix}_notes", ""))


def _add_bullets(doc, title, items):
    doc.add_heading(title, level=1)
    if items:
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph("—")


@app.post("/api/memo/docx")
def memo_to_docx(req: MemoDownloadRequest):
    memo = req.memo
    doc = Document()

    doc.add_heading(f"{PROJECT_NAME} — Status Memo", level=0)

    _add_kv(doc, "Report Date", memo.get("report_date", ""))
    _add_kv(
        doc,
        "Reporting Period",
        f"{memo.get('reporting_period_start', '')} to {memo.get('reporting_period_end', '')}",
    )
    _add_kv(doc, "Current Phase", memo.get("current_phase", ""))
    _add_kv(doc, "Prepared By", memo.get("prepared_by", ""))

    doc.add_heading("Status Overview", level=1)
    _add_status_block(doc, "Schedule", memo, "schedule")
    _add_status_block(doc, "Scope", memo, "scope")
    _add_status_block(doc, "Resources", memo, "resources")
    _add_status_block(doc, "Vendor", memo, "vendor")

    _add_bullets(doc, "Completed", memo.get("completed_items", []))
    _add_bullets(doc, "In Progress", memo.get("current_items", []))
    _add_bullets(doc, "Upcoming", memo.get("upcoming_items", []))

    doc.add_heading("Risks", level=1)
    risks = memo.get("risks", []) or []
    if risks:
        for r in risks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(r.get("description", "—"))).bold = True
            doc.add_paragraph(
                f"Likelihood: {r.get('likelihood', '—')} | Impact: {r.get('impact', '—')}"
            )
            doc.add_paragraph(f"Mitigation: {r.get('mitigation', '—')}")
    else:
        doc.add_paragraph("—")

    _add_bullets(doc, "Assumptions", memo.get("assumptions", []))

    doc.add_heading("Issues", level=1)
    issues = memo.get("issues", []) or []
    if issues:
        for i in issues:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(i.get("description", "—"))).bold = True
            doc.add_paragraph(
                f"Raised by: {i.get('raised_by', '—')} | Status: {i.get('status', '—')}"
            )
            doc.add_paragraph(f"Next step: {i.get('next_step', '—')}")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Decisions", level=1)
    decisions = memo.get("decisions", []) or []
    if decisions:
        for d in decisions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(d.get("description", "—"))).bold = True
            doc.add_paragraph(
                f"Made by: {d.get('made_by', '—')} | Date: {d.get('date', '—')}"
            )
    else:
        doc.add_paragraph("—")

    _add_bullets(doc, "Notes", memo.get("notes", []))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        content=buf.read(),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": 'attachment; filename="project-status-memo.docx"'
        },
    )


app.mount("/", StaticFiles(directory="static", html=True), name="static")
